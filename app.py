#THIS IS WORKING WITH ALL FEATURES 
from flask import Flask, render_template, request, redirect, url_for, session, make_response, send_file, flash
import os
from werkzeug.utils import secure_filename
from docx import Document
import PyPDF2
import speech_recognition as sr
from pydub import AudioSegment
from moviepy import VideoFileClip
from langdetect import detect
from deep_translator import GoogleTranslator
from datetime import datetime
import sqlite3
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import json
import io
import whisper
import easyocr
import tempfile

# ---------------- Database Setup ---------------- #
def init_db():
    conn = sqlite3.connect('smartdoc_data.db')
    cur = conn.cursor()
    cur.execute('''
    CREATE TABLE IF NOT EXISTS history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        file_name TEXT NOT NULL,
        file_type TEXT NOT NULL,
        analysis_date TEXT NOT NULL
    )
    ''')
    conn.commit()
    conn.close()

init_db()

def save_history(file_name, file_type):
    conn = sqlite3.connect('smartdoc_data.db')
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO history (file_name, file_type, analysis_date) VALUES (?, ?, ?)",
        (file_name, file_type, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    conn.commit()
    conn.close()

# ---------------- Flask Setup ---------------- #
app = Flask(__name__)
app.secret_key = 'sigma_secret_key'
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'png', 'jpg', 'jpeg', 'mp3', 'wav', 'mp4', 'mkv'}
MAX_SIZE_MB = 100

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_SIZE_MB * 1024 * 1024
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ---------------- Utility Functions ---------------- #
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text_for_xml(text):
    return ''.join(
        ch if ord(ch) in (0x09, 0x0A, 0x0D) or (0x20 <= ord(ch) <= 0xD7FF) or (0xE000 <= ord(ch) <= 0xFFFD)
        else ' '
        for ch in text
    )

def safe_delete(path):
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass

# ---------------- Audio & Video ---------------- #
def audio_to_text(audio_path, whisper_model_name="tiny"):
    tmp_wav = None
    try:
        if not audio_path.lower().endswith('.wav'):
            sound = AudioSegment.from_file(audio_path)
            sound = sound.set_frame_rate(16000).set_channels(1)
            tmp_wav = audio_path.rsplit('.', 1)[0] + "_converted.wav"
            sound.export(tmp_wav, format='wav')
            work_path = tmp_wav
        else:
            work_path = audio_path

        model = whisper.load_model(whisper_model_name)
        result = model.transcribe(work_path, verbose=False, fp16=False)
        return result.get("text", "").strip()

    except Exception as e:
        print("⚠️ Whisper error:", str(e))
        try:
            recognizer = sr.Recognizer()
            with sr.AudioFile(audio_path) as source:
                audio_data = recognizer.record(source)
                return recognizer.recognize_google(audio_data)
        except Exception:
            return "⚠️ Audio could not be transcribed."
    finally:
        if tmp_wav:
            safe_delete(tmp_wav)

def video_to_text(video_path):
    try:
        temp_audio = "temp_audio.wav"
        clip = VideoFileClip(video_path)
        clip.audio.write_audiofile(temp_audio, logger=None)

        recognizer = sr.Recognizer()

        # ✅ Step 1: Light calibration using short segment
        with sr.AudioFile(temp_audio) as source:
            recognizer.adjust_for_ambient_noise(source, duration=0.2)

        # ✅ Step 2: Reopen and record full audio (nothing skipped)
        with sr.AudioFile(temp_audio) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)

        return text.strip()

    except Exception as e:
        return f"❌ Video processing failed: {str(e)}"


# ---------------- DOCX Output ---------------- #
def save_to_docx_file(filename, text):
    base = os.path.splitext(filename)[0]
    doc = Document()
    doc.add_heading(f'Extracted Content: {filename}', 0)
    for line in text.splitlines():
        doc.add_paragraph(line)
    path = os.path.join(OUTPUT_FOLDER, f"{base}.docx")
    doc.save(path)

# ---------------- EasyOCR Extraction ---------------- #
reader_pool = {}

def get_easyocr_reader(lang_hint=None):
    if lang_hint == 'ta':
        key = 'ta'; langs = ['ta', 'en']
    elif lang_hint == 'kn':
        key = 'kn'; langs = ['kn', 'en']
    elif lang_hint in ['hi', 'mr']:
        key = 'hi_mr'; langs = ['en', 'hi', 'mr']
    else:
        key = 'en'; langs = ['en']

    if key not in reader_pool:
        reader_pool[key] = easyocr.Reader(langs, gpu=False)
    return reader_pool[key]

def extract_text_easyocr(image_path):
    try:
        default_reader = get_easyocr_reader('hi')
        text_lines = default_reader.readtext(image_path, detail=0)
        extracted_text = "\n".join(text_lines)

        if len(extracted_text.strip()) < 10:
            for hint in ['ta', 'kn', 'en']:
                reader = get_easyocr_reader(hint)
                text_lines = reader.readtext(image_path, detail=0)
                extracted_text = "\n".join(text_lines)
                if len(extracted_text.strip()) > 10:
                    break

        if not extracted_text.strip():
            return "⚠️ No readable text found in image."

        try:
            detected_lang = detect(extracted_text)
            if detected_lang != 'en':
                translated = GoogleTranslator(source='auto', target='en').translate(extracted_text)
                return f"🌍 Detected Language: {detected_lang.upper()}\n\n📝 Translated Text:\n{translated}"
        except:
            pass

        return extracted_text

    except Exception as e:
        return f"❌ OCR extraction failed: {str(e)}"

# ---------------- File Extractor ---------------- #
def extract_text_only(file_path, ext):
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().rstrip()
        elif ext == 'docx':
            doc = Document(file_path)
            return '\n'.join([p.text for p in doc.paragraphs])
        elif ext == 'pdf':
            text = ""
            with open(file_path, 'rb') as f:
                reader_pdf = PyPDF2.PdfReader(f)
                for page in reader_pdf.pages:
                    text += clean_text_for_xml(page.extract_text() or '') + '\n\n'
            return text
        elif ext in {'png', 'jpg', 'jpeg'}:
            return extract_text_easyocr(file_path)
        elif ext in {'mp3', 'wav'}:
            return audio_to_text(file_path)
        elif ext in {'mp4', 'mkv'}:
            return video_to_text(file_path)
        else:
            return "❌ Unsupported file type."
    except Exception as e:
        return f"❌ Extraction failed: {str(e)}"

# ---------------- Routes ---------------- #
@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST':
        session.clear()
        uploaded_files = request.files.getlist('files')
        output_lang = request.form.get('lang', 'original')
        errors = []
        result_keys = []

        if not uploaded_files or uploaded_files[0].filename == '':
            session['errors'] = ["⚠️ No files selected."]
            return redirect(url_for('results'))

        lock = threading.Lock()
        with ThreadPoolExecutor(max_workers=4) as executor:
            future_to_file = {}
            for idx, file in enumerate(uploaded_files):
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    ext = filename.rsplit('.', 1)[1].lower()
                    unique_name = f"{os.path.splitext(filename)[0]}_{idx}_{int(datetime.now().timestamp())}.{ext}"
                    file_path = os.path.join(UPLOAD_FOLDER, unique_name)
                    file.save(file_path)
                    future = executor.submit(extract_text_only, file_path, ext)
                    future_to_file[future] = (unique_name, ext)
                else:
                    errors.append(f"❌ {file.filename} not allowed.")

            for future in as_completed(future_to_file):
                unique_name, ext = future_to_file[future]
                raw_text = future.result()
                text = raw_text
                if output_lang != 'original':
                    try:
                        detected_lang = detect(raw_text)
                        if detected_lang != output_lang:
                            text = GoogleTranslator(source='auto', target=output_lang).translate(raw_text)
                    except Exception as e:
                        text += f"\n⚠️ Translation Error: {str(e)}"

                save_to_docx_file(unique_name, text)
                save_history(unique_name, ext)

                result_obj = {
                    'filename': unique_name,
                    'content': text,
                    'timestamp': datetime.now().strftime('%d %b %Y, %I:%M %p')
                }
                json_path = os.path.join(OUTPUT_FOLDER, f"{unique_name}.json")
                with open(json_path, 'w', encoding='utf-8') as jf:
                    json.dump(result_obj, jf, ensure_ascii=False, indent=2)
                with lock:
                    result_keys.append(unique_name)

        session['file_keys'] = result_keys
        session['errors'] = errors
        session.modified = True
        return redirect(url_for('results'))

    conn = sqlite3.connect('smartdoc_data.db')
    cur = conn.cursor()
    cur.execute("SELECT file_name, file_type, analysis_date FROM history ORDER BY analysis_date DESC LIMIT 5")
    recent_files = cur.fetchall()
    conn.close()
    return render_template('upload.html', recent_files=recent_files)

@app.route('/results')
def results():
    result_keys = session.get('file_keys', [])
    errors = session.get('errors', [])
    file_contents = []

    for key in result_keys:
        json_file = os.path.join(OUTPUT_FOLDER, f"{key}.json")
        if os.path.exists(json_file):
            with open(json_file, 'r', encoding='utf-8') as f:
                file_contents.append(json.load(f))

    if not file_contents and not errors:
        errors = ["⚠️ No files were processed. Please try again."]

    response = make_response(render_template('results.html', file_contents=file_contents, errors=errors))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    return response

@app.route('/download_text/<filename>')
def download_text(filename):
    json_path = os.path.join(OUTPUT_FOLDER, f"{filename}.json")
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            content = json.load(f)['content']
        buffer = io.BytesIO()
        buffer.write(content.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f"{filename}_SmartDoc.txt", mimetype='text/plain')
    else:
        flash("File not found for download.")
        return redirect(url_for('results'))

@app.errorhandler(413)
def too_large(e):
    return f"File too large. Maximum allowed size is {MAX_SIZE_MB}MB.", 413

@app.errorhandler(500)
def server_error(e):
    return "⚠️ Something went wrong! Try uploading fewer or smaller files.", 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=False)
